"""
Document parser - extracts text from various file formats
"""

import os
from pathlib import Path
from pypdf import PdfReader
from docx import Document


def extract_text_from_file(file_path):
    """
    Extract text from a file based on its extension
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx, .doc)
    - Text (.txt, .md)
    
    Args:
        file_path: Path to the file (string or Path object)
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file doesn't exist
    """
    
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = file_path.suffix.lower()
    
    # PDF files
    if extension == '.pdf':
        return extract_from_pdf(file_path)
    
    # Word documents
    elif extension in ['.docx', '.doc']:
        return extract_from_docx(file_path)
    
    # Text files
    elif extension in ['.txt', '.md', '.text']:
        return extract_from_text(file_path)
    
    else:
        raise ValueError(
            f"Unsupported file format: {extension}\n"
            f"Supported formats: .pdf, .docx, .doc, .txt, .md"
        )


def extract_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        reader = PdfReader(file_path)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(page_text)
        
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise ValueError("PDF appears to be empty or contains only images")
        
        return full_text
    
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")


def extract_from_docx(file_path):
    """Extract text from Word document"""
    try:
        doc = Document(file_path)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise ValueError("Word document appears to be empty")
        
        return full_text
    
    except Exception as e:
        raise ValueError(f"Error reading Word document: {str(e)}")


def extract_from_text(file_path):
    """Extract text from plain text file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                if text.strip():
                    return text
            
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any common encoding")
    
    except Exception as e:
        raise ValueError(f"Error reading text file: {str(e)}")


def get_file_info(file_path):
    """Get basic information about a file"""
    file_path = Path(file_path)
    
    return {
        "name": file_path.name,
        "size_kb": file_path.stat().st_size / 1024,
        "extension": file_path.suffix.lower(),
        "exists": file_path.exists()
    }
